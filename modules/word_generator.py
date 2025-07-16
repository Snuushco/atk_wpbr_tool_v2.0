from docx import Document
import os
import logging

def generate_word_from_template(data: dict, template_path: str, output_path: str):
    doc = Document(template_path)
    
    # Functie om placeholders in tekst te vervangen
    def replace_placeholders(text: str) -> str:
        if not text:
            return text
        result = text
        for key, value in data.items():
            placeholder = f"{{{{{key}}}}}"
            if placeholder in result:
                result = result.replace(placeholder, str(value))
        return result
    
    # Verwerk alle paragrafen
    for p in doc.paragraphs:
        if p.text:
            p.text = replace_placeholders(p.text)
        # Verwerk ook inline runs voor behoud van formatting
        for run in p.runs:
            if run.text:
                run.text = replace_placeholders(run.text)
    
    # Verwerk alle tabellen
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # Verwerk paragrafen in cellen
                for p in cell.paragraphs:
                    if p.text:
                        p.text = replace_placeholders(p.text)
                    # Verwerk ook inline runs in cellen
                    for run in p.runs:
                        if run.text:
                            run.text = replace_placeholders(run.text)
    
    # Verwerk headers
    for section in doc.sections:
        for header in section.header.paragraphs:
            if header.text:
                header.text = replace_placeholders(header.text)
            for run in header.runs:
                if run.text:
                    run.text = replace_placeholders(run.text)
    
    # Verwerk footers
    for section in doc.sections:
        for footer in section.footer.paragraphs:
            if footer.text:
                footer.text = replace_placeholders(footer.text)
            for run in footer.runs:
                if run.text:
                    run.text = replace_placeholders(run.text)
    
    # Log welke placeholders zijn vervangen
    logging.info(f"Word template verwerkt met data: {data}")
    
    # Sla het document op
    doc.save(output_path)
    return output_path 